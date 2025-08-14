#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将多个docx文件转换为可分片的HTML片段
用于个人网站子页面内容更新
"""

import os
import argparse
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.oxml.shared import qn
from docx.oxml import parse_xml
import html
import math

class DocxToHtmlConverter:
    def __init__(self, input_dir, output_dir, chunk_size=2000):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.images_dir = self.output_dir / "images"
        self.chunk_size = chunk_size
        
        # 文件映射关系
        self.file_mapping = {
            "产品需求文档.docx": "揭示产品思维",
            "产品分析报告.docx": "揭示产品思维", 
            "AI产品体验报告.docx": "洞察分析能力",
            "市场调研报告.docx": "洞察分析能力",
            "vibe开发总结.docx": "掌握工具应用",
            "学习总结与心得.docx": "沉淀学习心得"
        }
        
        # 板块内容存储
        self.section_contents = {}
        
        # 确保输出目录存在
        self.output_dir.mkdir(exist_ok=True)
        self.images_dir.mkdir(exist_ok=True)
    
    def extract_text_with_formatting(self, docx_path):
        """从docx文件中提取带格式的文本内容"""
        try:
            doc = Document(docx_path)
            content = []
            
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    continue
                
                # 处理段落格式
                para_html = self._process_paragraph(paragraph)
                if para_html:
                    content.append(para_html)
            
            # 处理表格
            for table in doc.tables:
                table_html = self._process_table(table)
                if table_html:
                    content.append(table_html)
            
            return "\n".join(content)
            
        except Exception as e:
            print(f"读取文件 {docx_path} 时出错: {e}")
            return f"<p>读取文档时出错: {e}</p>"
    
    def _process_paragraph(self, paragraph):
        """处理段落，保留格式"""
        if not paragraph.text.strip():
            return ""
        
        # 检查是否是标题
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
            level = min(max(level, 1), 6)  # 限制在h1-h6范围内
            return f"<h{level}>{self._process_runs(paragraph.runs)}</h{level}>"
        
        # 普通段落
        para_html = f"<p>{self._process_runs(paragraph.runs)}</p>"
        
        # 处理段落对齐
        if paragraph.alignment:
            align_map = {0: 'left', 1: 'center', 2: 'right', 3: 'justify'}
            if paragraph.alignment in align_map:
                para_html = para_html.replace('<p>', f'<p style="text-align: {align_map[paragraph.alignment]}">')
        
        return para_html
    
    def _process_runs(self, runs):
        """处理文本运行，保留加粗、斜体等格式"""
        if not runs:
            return ""
        
        html_parts = []
        for run in runs:
            text = html.escape(run.text)
            if not text.strip():
                continue
            
            # 构建HTML标签
            tags = []
            if run.bold:
                tags.append('strong')
            if run.italic:
                tags.append('em')
            if run.underline:
                tags.append('u')
            
            # 应用标签
            result = text
            for tag in tags:
                result = f"<{tag}>{result}</{tag}>"
            
            html_parts.append(result)
        
        return "".join(html_parts)
    
    def _process_table(self, table):
        """处理表格"""
        if not table.rows:
            return ""
        
        table_html = '<table border="1" style="border-collapse: collapse; width: 100%; margin: 20px 0; font-size: 14px;">'
        
        for i, row in enumerate(table.rows):
            # 判断是否是表头行
            tag = 'th' if i == 0 else 'td'
            table_html += '<tr>'
            
            for cell in row.cells:
                cell_text = html.escape(cell.text.strip())
                table_html += f'<{tag} style="padding: 8px; border: 1px solid #ddd; text-align: left;">{cell_text}</{tag}>'
            
            table_html += '</tr>'
        
        table_html += '</table>'
        return table_html
    
    def extract_images(self, docx_path):
        """从docx文件中提取图片"""
        try:
            doc = Document(docx_path)
            image_count = 0
            image_refs = []
            
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
                    file_name = os.path.basename(rel.target_ref)
                    name, ext = os.path.splitext(file_name)
                    
                    # 生成唯一文件名
                    image_name = f"{Path(docx_path).stem}_image_{image_count}{ext}"
                    image_path = self.images_dir / image_name
                    
                    # 保存图片
                    with open(image_path, 'wb') as f:
                        f.write(rel.target_part.blob)
                    
                    # 记录图片引用
                    image_refs.append({
                        'path': f"images/{image_name}",
                        'name': image_name
                    })
                    
                    print(f"提取图片: {image_path}")
            
            return image_refs
            
        except Exception as e:
            print(f"提取图片时出错: {e}")
            return []
    
    def add_images_to_content(self, content, image_refs):
        """在内容中添加图片引用"""
        if not image_refs:
            return content
        
        # 在内容末尾添加图片展示
        image_html = '<div class="document-images" style="margin-top: 30px; padding: 20px; background: #f9f9f9; border-radius: 8px;">'
        image_html += '<h3>文档图片</h3>'
        
        for img in image_refs:
            image_html += f'''
            <div style="margin: 15px 0; text-align: center;">
                <img src="{img['path']}" alt="{img['name']}" style="max-width: 100%; height: auto; border-radius: 4px; box-shadow: 0 2px 8px rgba(0,0,0,0.1);">
                <p style="margin-top: 8px; color: #666; font-size: 12px;">{img['name']}</p>
            </div>
            '''
        
        image_html += '</div>'
        return content + image_html
    
    def split_html_content(self, content, section_name):
        """将HTML内容按字符数分片"""
        # 移除HTML标签计算纯文本长度
        text_content = re.sub(r'<[^>]+>', '', content)
        total_chars = len(text_content)
        
        if total_chars <= self.chunk_size:
            return [content]
        
        # 计算分片数量
        num_chunks = math.ceil(total_chars / self.chunk_size)
        chunks = []
        
        # 按段落分片，确保不破坏HTML结构
        paragraphs = re.split(r'(<p[^>]*>.*?</p>|<h[1-6][^>]*>.*?</h[1-6]>|<table[^>]*>.*?</table>)', content, flags=re.DOTALL)
        
        current_chunk = ""
        current_chars = 0
        
        for para in paragraphs:
            if not para.strip():
                continue
            
            # 计算当前段落的纯文本长度
            para_text = re.sub(r'<[^>]+>', '', para)
            para_chars = len(para_text)
            
            # 如果添加当前段落会超过限制，且当前分片不为空，则保存当前分片
            if current_chunk and (current_chars + para_chars > self.chunk_size):
                chunks.append(current_chunk.strip())
                current_chunk = para
                current_chars = para_chars
            else:
                current_chunk += para
                current_chars += para_chars
        
        # 添加最后一个分片
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def save_chunks(self, section_name, chunks):
        """保存分片内容"""
        for i, chunk in enumerate(chunks, 1):
            filename = f"{section_name}_part{i}.html"
            filepath = self.output_dir / filename
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(f"<!-- {section_name} 第{i}部分 -->\n")
                f.write(chunk)
            
            print(f"已保存: {filename}")
    
    def process_files(self):
        """处理所有文件"""
        print(f"开始处理文档，输入目录: {self.input_dir}")
        print(f"输出目录: {self.output_dir}")
        print(f"分片大小: {self.chunk_size} 字符")
        print("-" * 50)
        
        # 按板块分组处理文件
        for docx_file, section_name in self.file_mapping.items():
            docx_path = self.input_dir / docx_file
            
            if not docx_path.exists():
                print(f"文件不存在: {docx_file}")
                continue
            
            print(f"\n处理文件: {docx_file} -> {section_name}")
            
            # 提取文本内容
            content = self.extract_text_with_formatting(docx_path)
            
            # 提取图片
            image_refs = self.extract_images(docx_path)
            
            # 添加图片到内容
            if image_refs:
                content = self.add_images_to_content(content, image_refs)
            
            # 将内容添加到对应板块
            if section_name not in self.section_contents:
                self.section_contents[section_name] = []
            
            self.section_contents[section_name].append({
                'source': docx_file,
                'content': content
            })
        
        # 合并同一板块的多个文档内容
        for section_name, docs in self.section_contents.items():
            print(f"\n合并板块: {section_name}")
            
            # 合并多个文档内容
            if len(docs) > 1:
                # 多个文档，添加文档标题分隔
                merged_content = ""
                for doc in docs:
                    merged_content += f'<h2>{doc["source"].replace(".docx", "")}</h2>\n'
                    merged_content += doc["content"] + "\n\n"
            else:
                # 单个文档
                merged_content = docs[0]["content"]
            
            # 分片处理
            chunks = self.split_html_content(merged_content, section_name)
            
            # 保存分片
            self.save_chunks(section_name, chunks)
        
        print(f"\n处理完成！共处理 {len(self.file_mapping)} 个文档")
        print(f"输出目录: {self.output_dir}")
        if any(self.section_contents.values()):
            print(f"图片目录: {self.images_dir}")

def main():
    parser = argparse.ArgumentParser(description='将docx文件转换为可分片的HTML片段')
    parser.add_argument('--input', '-i', default='.', help='输入文件夹路径 (默认: 当前目录)')
    parser.add_argument('--output', '-o', default='output', help='输出文件夹路径 (默认: output)')
    parser.add_argument('--chunk-size', '-c', type=int, default=2000, help='每片字符数 (默认: 2000)')
    
    args = parser.parse_args()
    
    # 检查输入目录
    if not os.path.exists(args.input):
        print(f"错误: 输入目录不存在: {args.input}")
        return
    
    # 创建转换器并处理文件
    converter = DocxToHtmlConverter(args.input, args.output, args.chunk_size)
    converter.process_files()

if __name__ == "__main__":
    main()
