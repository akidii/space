#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
读取.docx文件内容并转换为HTML格式的脚本
用于更新个人网站子页面内容
"""

import os
from docx import Document
from docx.shared import Inches
import re

def extract_text_from_docx(docx_path):
    """从.docx文件中提取文本内容"""
    try:
        doc = Document(docx_path)
        content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # 检查是否是标题
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                    content.append(f"<h{level}>{paragraph.text}</h{level}>")
                else:
                    content.append(f"<p>{paragraph.text}</p>")
        
        # 处理表格
        for table in doc.tables:
            table_html = "<table border='1' style='border-collapse: collapse; width: 100%; margin: 20px 0;'>"
            for row in table.rows:
                table_html += "<tr>"
                for cell in row.cells:
                    table_html += f"<td style='padding: 8px; border: 1px solid #ddd;'>{cell.text}</td>"
                table_html += "</tr>"
            table_html += "</table>"
            content.append(table_html)
        
        return "\n".join(content)
    except Exception as e:
        print(f"读取文件 {docx_path} 时出错: {e}")
        return f"<p>读取文档时出错: {e}</p>"

def extract_images_from_docx(docx_path, output_dir="images"):
    """从.docx文件中提取图片"""
    try:
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        doc = Document(docx_path)
        image_count = 0
        
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                image_name = f"{os.path.splitext(os.path.basename(docx_path))[0]}_image_{image_count}.png"
                image_path = os.path.join(output_dir, image_name)
                
                with open(image_path, 'wb') as f:
                    f.write(rel.target_part.blob)
                
                print(f"提取图片: {image_path}")
        
        return image_count
    except Exception as e:
        print(f"提取图片时出错: {e}")
        return 0

def main():
    """主函数"""
    print("开始读取.docx文件...")
    
    # 定义文件映射关系
    file_mapping = {
        "产品需求文档.docx": "product_requirements",
        "产品分析报告.docx": "product_analysis", 
        "AI产品体验报告.docx": "ai_experience",
        "市场调研报告.docx": "market_research",
        "vibe开发总结.docx": "vibe_summary",
        "学习总结与心得.docx": "learning_summary"
    }
    
    # 存储每个文档的内容
    document_contents = {}
    
    # 读取所有文档
    for docx_file, key in file_mapping.items():
        if os.path.exists(docx_file):
            print(f"正在读取: {docx_file}")
            content = extract_text_from_docx(docx_file)
            document_contents[key] = content
            
            # 尝试提取图片
            image_count = extract_images_from_docx(docx_file)
            if image_count > 0:
                print(f"从 {docx_file} 提取了 {image_count} 张图片")
        else:
            print(f"文件不存在: {docx_file}")
    
    # 生成合并后的内容
    merged_content = {
        "product": {
            "title": "揭示产品思维",
            "content": f"""
                <h2>产品需求文档</h2>
                {document_contents.get('product_requirements', '<p>文档内容读取失败</p>')}
                
                <h2>产品分析报告</h2>
                {document_contents.get('product_analysis', '<p>文档内容读取失败</p>')}
            """
        },
        "analysis": {
            "title": "洞察分析能力", 
            "content": f"""
                <h2>AI产品体验报告</h2>
                {document_contents.get('ai_experience', '<p>文档内容读取失败</p>')}
                
                <h2>市场调研报告</h2>
                {document_contents.get('market_research', '<p>文档内容读取失败</p>')}
            """
        },
        "tools": {
            "title": "掌握工具应用",
            "content": document_contents.get('vibe_summary', '<p>文档内容读取失败</p>')
        },
        "learning": {
            "title": "沉淀学习心得", 
            "content": document_contents.get('learning_summary', '<p>文档内容读取失败</p>')
        }
    }
    
    # 保存内容到文件
    for page_key, page_data in merged_content.items():
        output_file = f"extracted_content_{page_key}.html"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(f"<!-- {page_data['title']} 页面内容 -->\n")
            f.write(page_data['content'])
        print(f"已保存 {page_data['title']} 的内容到: {output_file}")
    
    print("\n所有文档内容提取完成！")
    return merged_content

if __name__ == "__main__":
    main()
